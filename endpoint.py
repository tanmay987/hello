from docx import Document
import io
from pydantic import BaseModel

# -------------------------
# 1) Extract text from docx
# -------------------------

def extract_docx(file):
    """
    Input: file (uploaded docx)
    Output: {"text": "..."}
    """
    content = file.read()
    doc = Document(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {"text": text}


# -------------------------
# 2) Build docx from text
# -------------------------

class BuildRequest(BaseModel):
    text: str

def build_docx(body: BuildRequest):
    """
    Input: {"text": "..."}
    Output: docx file
    """
    doc = Document()
    for line in body.text.split("\n"):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
